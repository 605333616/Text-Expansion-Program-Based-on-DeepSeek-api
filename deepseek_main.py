import threading
from openai import OpenAI
import tkinter as tk
from tkinter import *
from tkinter import filedialog, BOTTOM,ttk
from docx import Document
import sys

class PrintRedirector:
    def __init__(self, text_widget):
        self.text_space = text_widget
    def write(self, text):
        self.text_space.insert(tk.END, text)

def select_file():
    entryString = E1.get()
    system_string=E2.get("1.0","end")
    deepseek_modle=com.get()
    # 创建聊天完成请求
    client = OpenAI(
        api_key=entryString,  # 知识引擎原子能力 APIKey
        base_url="https://api.deepseek.com",
    )
    # 弹出文件选择对话框，并获取选中的文件路径
    selected_file_path = filedialog.askopenfilename()
    doc = Document(selected_file_path)
    # 初始化一个空列表来存储文字内容
    text_content = []
    word_content=[]
    # 遍历每个段落
    for paragraph in doc.paragraphs:
        # 获取段落文本
        content = paragraph.text
        # 将回车符号替换为空格
        content = content.replace('\r\n', ' ')
        text_content.append(content)
    # 构造 client
    # 流式
    # 请求
    for element in text_content:
        if element=='':
            pass
        else:
            reasoning_content = ""  # 定义完整思考过程
            answer_content = ""  # 定义完整回复
            is_answering = False  # 判断是否结束思考过程并开始回复
            chat_completion = client.chat.completions.create(
                model=deepseek_modle,
                messages=[
                    {
                        'role': 'system',
                        'content': system_string,
                    },
                    {
                        "role": "user",
                        "content": element
                    }
                ],
                stream=True,
            )

            textGather=''

            print("\n" + "=" * 20 + "思考过程" + "=" * 20 + "\n")
            output_text.yview_moveto(1.0)
            for chunk in chat_completion:
                # 处理usage信息
                if not getattr(chunk, 'choices', None):
                    print("\n" + "=" * 20 + "Token 使用情况" + "=" * 20 + "\n")
                    output_text.yview_moveto(1.0)
                    print(chunk.usage)
                    output_text.yview_moveto(1.0)
                    continue
                delta = chunk.choices[0].delta
                # 处理空内容情况
                if not getattr(delta, 'reasoning_content', None) and not getattr(delta, 'content', None):
                    continue
                # 处理开始回答的情况
                if not getattr(delta, 'reasoning_content', None) and not is_answering:
                    print("\n" + "=" * 20 + "完整回复" + "=" * 20 + "\n")
                    output_text.yview_moveto(1.0)
                    is_answering = True
                # 处理思考过程
                if getattr(delta, 'reasoning_content', None):
                    print(delta.reasoning_content, end='')
                    output_text.yview_moveto(1.0)
                    reasoning_content += delta.reasoning_content
                # 处理回复内容
                elif getattr(delta, 'content', None):
                    print(delta.content, end='')
                    textGather = textGather + delta.content
                    output_text.yview_moveto(1.0)
                    answer_content += delta.content



            word_content.append(textGather)



    full_text = '\n'.join(word_content)
    # 实例化一个Document对象，相当于打开word软件，新建一个空白文件
    doc = Document()
    # word文件尾部增加一个段落，并写入内容
    paragraph = doc.add_paragraph(full_text)
    # 保存word文件到当前文件夹
    address = selected_file_path.split('.', 1)
    doc.save(address[0] + '-已修改版本' + '.' + address[1])
    label["text"] = "生成完成"
    print("\n\n-----------生成完成--------------")
    output_text.yview_moveto(1.0)


class MyThread(threading.Thread):
    def __init__(self, func, *args):
        super().__init__()

        self.func = func
        self.args = args

        self.setDaemon(True)
        self.start()  # 在这里开始

    def run(self):
        self.func(*self.args)

# 创建Tkinter窗口
root = tk.Tk()
root.title("文本扩充器")
root.geometry("500x400")

L1 = Label(root, text="api密钥")
L1.place (x=15,y=18, width=60, height=30)
E1 = Entry(root, bd =5)
E1.place (x=70,y=20, width=400, height=30)

L2 = Label(root, text="系统提示词")
L2.place (x=8,y=53, width=60, height=30)
E2 = Text(root, bd =5)
E2.place (x=70,y=55, width=400, height=90)

#下拉选项框
xVariable = tk.StringVar()
com = ttk.Combobox(root, textvariable=xVariable)     # #创建下拉菜单
com.place (x=50,y=165, width=140, height=25)     # #将下拉菜单绑定到窗体
com["value"] = ("deepseek-chat", "deepseek-reasoner")
# 创建并配置按钮，点击时调用select_file函数
button = tk.Button(root, text="选择文件", command=lambda :MyThread(select_file))
button.place (x=200,y=160, width=90, height=30)

label = tk.Label(root, text="选择文件开始扩充文本",font=('Calibri 15 bold'))
label.place (x=45,y=300, width=400, height=50)

#输出框
output_text = tk.Text(root)
#滚动条
scrollbar = tk.Scrollbar(root)
scrollbar.place(x=495,y=200,width=5, height=80)
# 将滚动条与Text组件关联
output_text.config(yscrollcommand=scrollbar.set)
scrollbar.config(command=output_text.yview)

output_text.place(x=5,y=200,width=490, height=80)

sys.stdout = PrintRedirector(output_text)


statement=tk.Label(root, text="本项目完全免费，\n用户可以自由下载、使用以及分发本项目的代码和相关资源。\n你无需支付任何费用即可享受本项目的功能和服务。\n开发者不对因使用本项目而产生的任何直接或间接损失承担责任。\n用户在使用本项目时应自行承担风险，并对可能的后果负责。",
                   font=('Arial', 10))
statement.place (x=45,y=300, width=400, height=100)

root.resizable(0,0)
# 运行Tkinter事件循环
root.mainloop()
